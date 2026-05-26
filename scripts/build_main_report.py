from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
REPORT_PATH = OUTPUT_DIR / "Champo_Carpets_Main_Report.docx"


def load_json(name: str) -> dict:
    return json.loads((OUTPUT_DIR / name).read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))

    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(8)

    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(12)
        styles[style_name].paragraph_format.line_spacing = 1.5

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = setup_document()

    target = load_json("part1_sample_target_summary.json")
    lr = load_json("part3_logistic_metrics.json")["metrics"]
    tree_models = pd.read_csv(OUTPUT_DIR / "part4_tree_model_comparison.csv")
    lr_features = pd.read_csv(OUTPUT_DIR / "part3_logistic_top_coefficients.csv").head(8)
    rf_features = pd.read_csv(OUTPUT_DIR / "part4_random_forest_feature_importance.csv").head(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Main Report: Champo Carpets ML Analysis")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(12)

    add_heading(doc, "1. Problem Understanding")
    add_para(
        doc,
        "Champo Carpets operates in a B2B market where sample development is an important part of the sales process. "
        "The company sends carpet samples to potential customers, but sample design and production require time, material, and design effort. "
        "When samples do not convert into orders, the company incurs avoidable cost and loses sales productivity. "
        "The core business problem is therefore to identify which sample requests are more likely to convert into actual orders and which combinations need closer review before resources are committed.",
    )
    add_para(
        doc,
        f"The sample-only data contained {target['rows']} records and showed an overall conversion rate of approximately {target['conversion_rate']:.1%}. "
        "This indicates that only about one in five samples resulted in an order. A data-driven approach can help Champo prioritize high-potential samples, improve customer targeting, and reduce inefficient sample creation.",
    )

    add_heading(doc, "2. Feature Engineering And Data Preparation")
    add_para(
        doc,
        "The analysis used the provided workbook tabs in stages. The full order and sample sheet was used to understand the complete transaction base, the order-only sheet was used for customer and revenue context, and the sample-only sheet was used for model development because it contained the target variable, Order Conversion.",
    )
    add_bullet(doc, "The target variable was Order Conversion, coded as 1 for converted samples and 0 for non-converted samples.")
    add_bullet(doc, "Categorical fields such as CustomerCode, CountryName, ITEM_NAME, and ShapeName were encoded into model-ready indicator variables.")
    add_bullet(doc, "Numerical fields such as QtyRequired and AreaFt were cleaned and used directly.")
    add_bullet(doc, "Country dummy variables already present in the sheet were removed where CountryName was used, to reduce duplicate signals.")
    add_bullet(doc, "The data was split into 80% training data and 20% test data, keeping the conversion/non-conversion balance consistent across both sets.")

    add_heading(doc, "3. Model Development")
    add_para(
        doc,
        "Three model families were developed to address the prediction task. Logistic Regression was selected as the primary interpretable model because it estimates how each feature changes the probability of conversion. Decision Tree learning was used to capture rule-based patterns, while Random Forest was used as an advanced ensemble method to reduce the instability of a single tree and compare feature importance.",
    )
    add_para(
        doc,
        "The Logistic Regression model was trained on the encoded sample dataset and evaluated on the 20% holdout test set. Decision Tree and Random Forest models used the same train-test split so that model performance could be compared fairly. The modelling objective was not only accuracy, but also business usefulness: Champo needs a model that can identify promising opportunities while remaining understandable for sales and operations teams.",
    )

    add_heading(doc, "4. Model Validation And Diagnostics")
    add_para(
        doc,
        "Model performance was evaluated using accuracy, precision, recall, and F1 score. Accuracy shows overall correctness, precision shows how reliable positive conversion predictions are, recall shows how many actual conversions the model captures, and F1 score balances precision and recall. Because only about 20% of samples converted, F1 score is more informative than accuracy alone.",
    )
    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1 Score"],
        [
            ["Logistic Regression", lr["accuracy"], lr["precision"], lr["recall"], lr["f1"]],
            *[
                [row["model"], row["accuracy"], row["precision"], row["recall"], row["f1"]]
                for _, row in tree_models.iterrows()
            ],
        ],
    )
    add_para(
        doc,
        "Logistic Regression performed best overall based on F1 score and interpretability. It achieved 84.26% accuracy, 73.58% precision, 33.48% recall, and an F1 score of 46.02%. The recall value shows that the model still misses some actual conversions, but its precision indicates that predicted conversions are relatively reliable. Random Forest had higher precision but lower recall, making it more conservative in identifying conversions.",
    )

    add_heading(doc, "5. Interpretation Using Explainable AI Techniques")
    add_para(
        doc,
        "Explainability was handled through Logistic Regression coefficients and Random Forest feature importance. Logistic Regression coefficients show the direction and relative strength of each feature, while Random Forest split counts show which variables were repeatedly useful in separating converted and non-converted samples.",
    )
    add_table(
        doc,
        ["Logistic Regression Feature", "Coefficient"],
        [[row["feature"], round(row["coefficient"], 4)] for _, row in lr_features.iterrows()],
    )
    add_table(
        doc,
        ["Random Forest Feature", "Importance Signal"],
        [[row["feature"], int(row["split_count"])] for _, row in rf_features.iterrows()],
    )
    add_para(
        doc,
        "Both approaches point to similar drivers. AreaFt and QtyRequired were important predictors, suggesting that larger-area and higher-quantity sample requests are more likely to become orders. Product type also mattered, with features such as Knotted and Other showing positive signals, while some categories such as Hand_Tufted and Double_Back showed weaker or negative influence in the Logistic Regression model.",
    )

    add_heading(doc, "6. Insights And Deployment Strategy")
    add_para(
        doc,
        "The analysis suggests that Champo should move from a purely reactive sample process to a scored and prioritized sales process. Each new sample request can be passed through the model to generate a conversion probability. Sales and design teams can then classify requests into high, medium, and low priority groups before committing resources.",
    )
    add_bullet(doc, "High-probability samples should receive faster design support and active sales follow-up.")
    add_bullet(doc, "Medium-probability samples should be reviewed using customer history, product type, and expected order value.")
    add_bullet(doc, "Low-probability samples should require stronger justification before costly sample production.")
    add_bullet(doc, "Customer segmentation should be used alongside the model so that high-value customers are not rejected only because of one low model score.")
    add_para(
        doc,
        "For deployment, the model can initially be used as a decision-support tool rather than an automated decision system. A simple scoring sheet or dashboard can show conversion probability, important drivers, and recommended action. Champo should retrain the model periodically as new ERP data becomes available.",
    )

    add_heading(doc, "7. Specific Recommendations Addressing The Case Questions")
    add_bullet(doc, "Use descriptive analytics to track conversion patterns by country, customer, product type, shape, quantity, and area.")
    add_bullet(doc, "Use Logistic Regression as the main explainable conversion model because it performed best overall and is easy to communicate to managers.")
    add_bullet(doc, "Use Decision Tree and Random Forest as supporting models to validate non-linear patterns and confirm important drivers.")
    add_bullet(doc, "Prioritize sample creation for combinations with higher predicted conversion probability, especially larger AreaFt and higher QtyRequired opportunities.")
    add_bullet(doc, "Reduce avoidable sample cost by pre-qualifying low-probability requests before design and production resources are committed.")
    add_bullet(doc, "Segment B2B customers using order value, area, quantity, and product mix so that sales strategy is tailored rather than uniform.")
    add_bullet(doc, "Use association and recommendation inputs for customer-specific design suggestions, especially for color and product combinations.")
    add_para(
        doc,
        "Overall, machine learning can create value for Champo by improving sample prioritization, reducing wasted design effort, increasing sales-team focus, and making the B2B sales process more evidence-based. The recommended starting point is an interpretable Logistic Regression scoring model supported by periodic comparison with tree-based models.",
    )

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
